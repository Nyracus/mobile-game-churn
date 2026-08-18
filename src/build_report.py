"""Build the final structured project report as a polished DOCX."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIGURES = REPORTS / "figures"
DELIVERABLES = ROOT / "deliverables"
FINAL_DOCX = DELIVERABLES / "mobile_game_churn_project_report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 55, 72)
MUTED = RGBColor(96, 106, 116)
LIGHT_FILL = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("Table widths must sum to 9360 DXA.")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_dxa)):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
        set_cell_shading(cell, PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = INK
            run.font.size = Pt(9)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.color.rgb = MUTED
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)


def add_figure(doc: Document, filename: str, caption: str, width=6.25) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(FIGURES / filename), width=Inches(width))
    add_caption(doc, caption)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(30, 35, 40)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    if "Lead" not in styles:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Lead"]
    lead.font.name = "Calibri"
    lead.font.size = Pt(12)
    lead.font.color.rgb = INK
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "EARLY PREDICTION OF PLAYER CHURN  |  CSE437 PROJECT REPORT"
    header.style = styles["Normal"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("MACHINE LEARNING PROJECT REPORT")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Early Prediction of Player Churn in Mobile Games")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("Gameplay event aggregation, comparative modeling, and explainable AI")
    run.font.size = Pt(15)
    run.font.color.rgb = DARK_BLUE

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run("CSE437  |  BRAC University  |  August 2026").font.color.rgb = MUTED

    doc.add_paragraph()
    lead = doc.add_paragraph(style="Lead")
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.add_run("Research question: Can early level-attempt behavior identify players who will become inactive in the following week?")
    doc.add_page_break()


def build_report() -> Path:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    metrics = pd.read_csv(REPORTS / "metrics.csv")
    selection = pd.read_csv(REPORTS / "model_selection_summary.csv")
    intervals = pd.read_csv(REPORTS / "metric_confidence_intervals.csv").set_index("metric")
    shap_imp = pd.read_csv(REPORTS / "shap_importance.csv")
    errors = pd.read_csv(REPORTS / "error_analysis_summary.csv")

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Project Goal", level=1)
    doc.add_paragraph(
        "This project develops an early-warning machine-learning system for mobile-game player churn. "
        "The goal is to predict whether a player will become inactive in the following week using only "
        "behavior observed during the initial gameplay window, compare several classification approaches, "
        "and explain the behavioral patterns behind the strongest model so that retention teams can identify "
        "where supportive interventions may be useful."
    )

    doc.add_heading("2. Method", level=1)
    doc.add_heading("2.1 Dataset and prediction target", level=2)
    doc.add_paragraph(
        "The public Prediction of User Loss in Mobile Games dataset contains anonymized player labels, "
        "level-attempt logs, and level metadata. A label of 1 indicates that the player did not log in during "
        "the following week; label 0 indicates retention. The source files use tab delimiters despite their "
        ".csv extension. The official train, development, and test memberships were preserved."
    )
    add_table(
        doc,
        ["Source", "Role", "Records"],
        [
            ["train.csv", "Model development labels", "8,158 players"],
            ["dev.csv", "Untouched final evaluation labels", "2,658 players"],
            ["test.csv", "Unlabeled final predictions", "2,773 players"],
            ["level_seq.csv", "Timestamped gameplay attempts", "2,194,351 events"],
            ["level_meta.csv", "Level difficulty and duration metadata", "1,509 levels"],
        ],
        [1800, 4860, 2700],
    )
    source = doc.add_paragraph(style="Caption")
    source.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source.add_run("Dataset source: Kaggle, manchvictor/prediction-of-user-loss-in-mobile-games.")

    doc.add_heading("2.2 Data cleaning and integration", level=2)
    doc.add_paragraph(
        "The pipeline parsed the five tab-separated sources with explicit date-time conversion, checked player "
        "identifiers for uniqueness within each official split, joined level metadata to every attempt using "
        "level_id, sorted attempts by player and time, and rejected any split member without gameplay events. "
        "Generated infinite values were converted to missing values and handled by median imputation inside each "
        "model pipeline. This fold-local preprocessing prevents validation information from leaking into training."
    )

    doc.add_heading("2.3 Player-level feature engineering", level=2)
    doc.add_paragraph(
        "The 2.19 million event rows were aggregated into 41 numeric predictors per player. The feature groups "
        "cover activity volume (attempts, unique levels, sessions, active days), progression (maximum and mean "
        "level), performance (success, failure, retry, help use, remaining steps), duration, temporal behavior "
        "(weekend/night activity, observation span, recency), level difficulty, and recent-form changes computed "
        "over the last 25 attempts. user_id was retained only for joins and never entered a model."
    )

    doc.add_heading("2.4 Class balance and evaluation criteria", level=2)
    doc.add_paragraph(
        "Churned players form 33.5% of the training set, placing the project in Group 2: mild imbalance (minority "
        "class between 5% and 40%). In accordance with the project guidance, model comparison emphasizes AUROC, "
        "precision, and recall; F1-score is also reported for direct comparison. AUPRC is retained as an additional "
        "minority-class summary. Accuracy is shown but is not used as the primary selection criterion."
    )
    add_figure(doc, "eda_class_distribution.png", "Figure 1. Training-set class distribution; the minority churn class represents 33.5% of players.", 5.4)

    doc.add_heading("2.5 Research workflow", level=2)
    add_figure(doc, "research_workflow.png", "Figure 2. End-to-end research workflow with the development split isolated from all selection decisions.", 6.45)

    doc.add_heading("2.6 Model training and selection", level=2)
    doc.add_paragraph(
        "Four model families were compared: Logistic Regression, Random Forest, Histogram Gradient Boosting, and "
        "XGBoost. Each family received 15 randomized hyperparameter configurations evaluated through five-fold "
        "stratified cross-validation on the training split. The configuration with the highest mean cross-validated "
        "AUROC defined the winner. The Random Forest used 770 trees, maximum depth 8, square-root feature sampling, "
        "and minimum leaf size 4."
    )

    doc.add_heading("2.7 Threshold selection, uncertainty, and explanation", level=2)
    doc.add_paragraph(
        "Cross-validated out-of-fold probabilities from the selected training model were used to choose a decision "
        "threshold that maximized churn F1, with recall used to break ties. The selected threshold was 0.320. Only "
        "after fixing the model and threshold was the official development set evaluated once. Uncertainty was "
        "estimated using 1,000 bootstrap resamples. SHAP values summarized global feature influence and an individual "
        "prediction, while false-positive and false-negative groups supported error analysis."
    )
    add_figure(doc, "threshold_selection.png", "Figure 3. Training out-of-fold threshold selection; no development labels were used.", 5.9)

    doc.add_page_break()
    doc.add_heading("3. Results", level=1)
    doc.add_heading("3.1 Comparative model performance", level=2)
    comparison_rows = []
    for row in selection.itertuples():
        comparison_rows.append(
            [
                row.model.replace("_", " ").title(),
                f"{row.cv_roc_auc:.4f}",
                f"{row.dev_roc_auc:.4f}",
                f"{row.dev_precision:.4f}",
                f"{row.dev_recall:.4f}",
                f"{row.dev_f1:.4f}",
                f"{row.dev_pr_auc:.4f}",
            ]
        )
    table = add_table(
        doc,
        ["Model", "CV AUROC", "Dev AUROC", "Precision", "Recall", "F1", "AUPRC"],
        comparison_rows,
        [2160, 1200, 1200, 1200, 1200, 1200, 1200],
    )
    for cell in table.rows[1].cells:
        set_cell_shading(cell, "EDF6FA")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    add_caption(doc, "Table 2. Training-CV selection evidence and final development-set metrics for all four model families.")
    add_figure(doc, "tuned_model_comparison.png", "Figure 4. Comparative performance across model families.", 6.25)

    doc.add_heading("3.2 Final Random Forest evaluation", level=2)
    best = metrics.iloc[0]
    final_rows = []
    for metric, label in (
        ("dev_roc_auc", "AUROC"),
        ("dev_pr_auc", "AUPRC"),
        ("dev_precision", "Precision"),
        ("dev_recall", "Recall"),
        ("dev_f1", "F1-score"),
        ("dev_accuracy", "Accuracy"),
    ):
        key = metric.removeprefix("dev_")
        if key in intervals.index:
            ci = f"{intervals.loc[key, 'lower_95']:.4f}-{intervals.loc[key, 'upper_95']:.4f}"
        else:
            ci = "Not bootstrapped"
        final_rows.append([label, f"{best[metric]:.4f}", ci])
    add_table(doc, ["Metric", "Development result", "Bootstrap 95% CI"], final_rows, [2700, 2700, 3960])
    add_caption(doc, "Table 3. Final holdout performance using the training-only threshold of 0.320.")
    doc.add_paragraph(
        "The model identified 704 of 901 churned development players, producing recall of 0.7814. Its precision of "
        "0.5789 reflects 512 retained users flagged as at risk. The AUROC confidence interval (0.7860-0.8196) and "
        "similar training-CV performance indicate limited split-specific degradation, although external validation "
        "is still needed."
    )
    add_figure(doc, "confusion_matrix.png", "Figure 5. Development-set confusion matrix at the 0.320 decision threshold.", 4.9)
    add_figure(doc, "precision_recall_curve.png", "Figure 6. Precision-recall curve on the untouched development set.", 5.4)
    add_figure(doc, "roc_curve.png", "Figure 7. ROC curve on the untouched development set.", 5.4)

    doc.add_heading("3.3 Explainability and behavioral findings", level=2)
    doc.add_paragraph(
        "SHAP analysis ranked inactivity recency as the strongest model signal, followed by weekend concentration, "
        "observation span, active days, total duration, and session count. These results describe associations within "
        "this fitted model; they do not prove that changing any one behavior will cause retention."
    )
    shap_rows = [
        [row.feature.replace("_", " ").title(), f"{row.mean_absolute_shap:.4f}"]
        for row in shap_imp.head(10).itertuples()
    ]
    add_table(doc, ["Predictor", "Mean absolute SHAP"], shap_rows, [6660, 2700])
    add_caption(doc, "Table 4. Ten most influential features in the selected Random Forest.")
    add_figure(doc, "shap_beeswarm.png", "Figure 8. SHAP distribution showing feature direction and magnitude for development players.", 6.25)

    doc.add_page_break()
    doc.add_heading("3.4 Error analysis", level=2)
    error_rows = []
    for row in errors.itertuples():
        error_rows.append(
            [
                row.outcome.replace("_", " ").title(),
                f"{int(row.players):,}",
                f"{row.mean_probability:.3f}",
                f"{row.median_hours_since_last_activity:.1f}",
                f"{row.median_active_days:.1f}",
                f"{row.median_session_count:.1f}",
            ]
        )
    add_table(
        doc,
        ["Outcome", "Players", "Mean risk", "Median recency (h)", "Active days", "Sessions"],
        error_rows,
        [2400, 1050, 1350, 1860, 1350, 1350],
    )
    add_caption(doc, "Table 5. Behavioral profile of correct and incorrect development predictions.")
    doc.add_paragraph(
        "False negatives looked more like retained users: their median recency was only 5.4 hours, with four active "
        "days and eight sessions. False positives showed the opposite pattern, with median recency of 32.0 hours and "
        "only four sessions. This overlap explains why behavior-based churn prediction remains probabilistic and "
        "suggests that any intervention should be low-cost and reversible."
    )

    doc.add_heading("4. Conclusion", level=1)
    doc.add_paragraph(
        "Early gameplay logs contain sufficient signal to support useful, though imperfect, churn-risk ranking. "
        "Among four tuned classifiers, Random Forest achieved the strongest training cross-validated AUROC and "
        "generalized to a development AUROC of 0.8024, recall of 0.7814, and F1-score of 0.6651. Recency, continuity "
        "of activity, and progression intensity were the most influential predictors. The model can guide supportive "
        "retention experiments, but deployment should wait for validation on another time period or game, probability "
        "calibration monitoring, and controlled testing to distinguish predictive association from causal impact."
    )

    doc.add_heading("References", level=1)
    doc.add_paragraph(
        "ManVictor. Prediction of User Loss in Mobile Games. Kaggle dataset. "
        "https://www.kaggle.com/datasets/manchvictor/prediction-of-user-loss-in-mobile-games"
    )
    doc.add_paragraph(
        "Lundberg, S. M., & Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. "
        "Advances in Neural Information Processing Systems, 30."
    )

    doc.save(FINAL_DOCX)
    print(f"Saved {FINAL_DOCX}")
    return FINAL_DOCX


if __name__ == "__main__":
    build_report()
